# md_to_docx.py — simple markdown-ish to docx
from pathlib import Path
import re
from docx import Document
from docx.shared import Pt, Cm

def convert(md_path: Path, out_path: Path) -> None:
    md = md_path.read_text(encoding="utf-8")
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2)
        s.bottom_margin = Cm(2)
        s.left_margin = Cm(2.2)
        s.right_margin = Cm(2.2)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)
    lines = md.splitlines()
    i = 0

    def flush_table(buf):
        rows = []
        for row in buf:
            if re.match(r"^\|?\s*-+", row):
                continue
            rows.append([c.strip() for c in row.strip().strip("|").split("|")])
        if not rows:
            return
        cols = max(len(r) for r in rows)
        t = doc.add_table(rows=len(rows), cols=cols)
        t.style = "Table Grid"
        for ri, r in enumerate(rows):
            for ci in range(cols):
                val = r[ci] if ci < len(r) else ""
                val = re.sub(r"\*\*(.+?)\*\*", r"\1", val)
                cell = t.rows[ri].cells[ci]
                cell.text = val
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9)
                        if ri == 0:
                            run.bold = True
        doc.add_paragraph("")

    while i < len(lines):
        line = lines[i]
        if line.strip().startswith("|"):
            buf = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                buf.append(lines[i])
                i += 1
            flush_table(buf)
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), 0)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), 1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), 2)
        elif line.strip() in ("---", ""):
            pass
        elif line.strip().startswith("```"):
            i += 1
            code = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code.append(lines[i])
                i += 1
            p = doc.add_paragraph("\n".join(code))
            for run in p.runs:
                run.font.name = "Consolas"
                run.font.size = Pt(9)
        elif line.strip().startswith(">"):
            p = doc.add_paragraph()
            p.add_run(line.lstrip("> ").strip()).italic = True
        elif line.strip().startswith("- "):
            doc.add_paragraph(
                re.sub(r"\*\*(.+?)\*\*", r"\1", line.strip()[2:]),
                style="List Bullet",
            )
        else:
            p = doc.add_paragraph()
            for part in re.split(r"(\*\*.+?\*\*)", line):
                if part.startswith("**") and part.endswith("**"):
                    p.add_run(part[2:-2]).bold = True
                else:
                    p.add_run(part)
        i += 1
    doc.save(out_path)
    print("OK", out_path, out_path.stat().st_size)


if __name__ == "__main__":
    convert(Path("report/RAPOR_Birliktelik.md"), Path("report/RAPOR_Birliktelik.docx"))
    # recreate classification docx if missing
    p = Path("report/RAPOR_Siniflandirma.md")
    if p.exists():
        convert(p, Path("report/RAPOR_Siniflandirma.docx"))
